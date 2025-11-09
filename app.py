"""Aplicação Tkinter para extrair itens de ATAs em PDF e preencher modelo de ofício Word."""
import copy
import os
import re
from datetime import datetime
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import tkinter as tk
from tkinter import filedialog, messagebox

try:
    import pdfplumber
except ImportError as exc:  # pragma: no cover - feedback para usuário
    raise SystemExit(
        "A biblioteca 'pdfplumber' é necessária. Instale-a com 'pip install pdfplumber'."
    ) from exc

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - feedback para usuário
    raise SystemExit(
        "A biblioteca 'python-docx' é necessária. Instale-a com 'pip install python-docx'."
    ) from exc


HEADER_HINTS: Dict[str, Tuple[str, ...]] = {
    "item": ("ITEM", "Nº", "N°", "NUM", "N. DO ITEM"),
    "description": ("DESCRI", "ESPECIF", "DESCRIÇÃO"),
    "quantity": ("QUANT", "QTD", "QTDE", "QUANTIDADE"),
    "unit_value": ("VALOR UNIT", "VLR UNIT", "UNITÁRIO", "UNIT"),
    "total_value": ("VALOR TOTAL", "TOTAL (R$)", "TOTAL"),
}


def normalize_text(value: Optional[str]) -> str:
    return (value or "").strip()


def normalize_item_code(value: str) -> str:
    """Normaliza código para comparação (apenas dígitos)."""
    digits = re.sub(r"\D", "", value or "")
    return digits or normalize_text(value).lower()


def find_header_indices(table: Sequence[Sequence[str]]) -> Optional[Tuple[int, Dict[str, int]]]:
    """Retorna índice da linha de cabeçalho e mapa de colunas relevantes."""
    for idx, row in enumerate(table):
        upper = [normalize_text(cell).upper() for cell in row]
        if not any(upper):
            continue
        mapping: Dict[str, int] = {}
        for col_index, cell in enumerate(upper):
            for key, hints in HEADER_HINTS.items():
                if key in mapping:
                    continue
                if any(hint in cell for hint in hints):
                    mapping[key] = col_index
        if all(k in mapping for k in ("item", "description", "quantity", "unit_value")):
            return idx, mapping
    return None


def extract_table_items(table: Sequence[Sequence[str]], header_idx: int, mapping: Dict[str, int]) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    current: Optional[Dict[str, str]] = None

    for raw_row in table[header_idx + 1 :]:
        row = [normalize_text(cell) for cell in raw_row]
        if not any(row):
            current = None
            continue

        def get(col_key: str) -> str:
            col = mapping.get(col_key)
            if col is None:
                return ""
            return row[col] if col < len(row) else ""

        item_code = get("item")
        desc_text = get("description")

        if item_code:
            current = {
                "item": item_code,
                "description": desc_text,
                "quantity": get("quantity"),
                "unit_value": get("unit_value"),
                "total_value": get("total_value"),
            }
            if any(current.values()):
                items.append(current)
        elif current and desc_text:
            # Linhas adicionais para a descrição
            current["description"] = (current["description"] + "\n" + desc_text).strip()
            if mapping.get("quantity") is not None and get("quantity"):
                current["quantity"] = get("quantity") or current["quantity"]
            if mapping.get("unit_value") is not None and get("unit_value"):
                current["unit_value"] = get("unit_value") or current["unit_value"]
            if mapping.get("total_value") is not None and get("total_value"):
                current["total_value"] = get("total_value") or current["total_value"]

    return items


def extract_items_from_pdf(path: str) -> List[Dict[str, str]]:
    try:
        with pdfplumber.open(path) as pdf:
            best_table: Optional[List[List[str]]] = None
            best_header: Optional[Tuple[int, Dict[str, int]]] = None
            best_length = 0

            for page in pdf.pages:
                for raw_table in page.extract_tables() or []:
                    table = [[cell or "" for cell in row] for row in raw_table if row]
                    if not table:
                        continue
                    header_info = find_header_indices(table)
                    if not header_info:
                        continue
                    header_idx, mapping = header_info
                    length = len(table) - (header_idx + 1)
                    if length <= 0:
                        continue
                    if best_table is None or length > best_length:
                        best_table = table
                        best_header = header_info
                        best_length = length
            if best_table and best_header:
                header_idx, mapping = best_header
                return extract_table_items(best_table, header_idx, mapping)
    except Exception as exc:  # pragma: no cover - gui feedback
        raise RuntimeError(f"Falha ao processar o PDF '{os.path.basename(path)}': {exc}") from exc
    return []


def parse_selected_items(text: str) -> List[str]:
    tokens = [normalize_text(token) for token in text.split(",")]
    return [token for token in tokens if token]


def filter_items(items: Iterable[Dict[str, str]], selected: Sequence[str]) -> List[Dict[str, str]]:
    if not selected:
        return list(items)

    normalized_targets = {normalize_item_code(value) for value in selected}
    filtered: List[Dict[str, str]] = []
    for item in items:
        code = normalize_text(item.get("item"))
        if not code:
            continue
        normalized_code = normalize_item_code(code)
        if normalized_code in normalized_targets or code in selected:
            filtered.append(item)
    return filtered


def replace_placeholders(doc: Document, replacements: Dict[str, str]) -> None:
    # Parágrafos principais
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    # Dentro de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)


def duplicate_row(table, row_idx: int):
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(tr)
    tbl.append(new_tr)
    return table.rows[-1]


def fill_items_table(doc: Document, items: Sequence[Dict[str, str]]) -> None:
    if not items:
        return

    target_table = None
    for table in doc.tables:
        header = " ".join(cell.text.upper() for cell in table.rows[0].cells)
        if all(keyword in header for keyword in ("ITEM", "QUANT")) and "VALOR" in header:
            target_table = table
            break
    if target_table is None:
        raise RuntimeError("Não foi possível localizar a tabela de itens no modelo Word.")

    if len(target_table.rows) < 2:
        duplicate_row(target_table, len(target_table.rows) - 1)

    template_row_idx = 1 if len(target_table.rows) > 1 else 0

    for idx, item in enumerate(items):
        if idx == 0:
            row = target_table.rows[template_row_idx]
        else:
            row = duplicate_row(target_table, template_row_idx)
        cells = row.cells
        mapping = {
            "item": 0,
            "description": 1,
            "quantity": 2,
            "unit_value": 3,
            "total_value": 4,
        }
        for key, cell_idx in mapping.items():
            if cell_idx >= len(cells):
                continue
            text_value = normalize_text(item.get(key, ""))
            cells[cell_idx].text = text_value


class AtaApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Gerador de Ofício a partir de ATA")
        self.pdf_paths: List[str] = []
        self.template_path: Optional[str] = None
        self._build_widgets()

    def _build_widgets(self) -> None:
        self.geometry("680x420")
        self.resizable(False, False)

        padding = {"padx": 10, "pady": 5}

        btn_pdf = tk.Button(self, text="Selecionar PDF da ATA", command=self.select_pdfs)
        btn_pdf.grid(row=0, column=0, sticky="w", **padding)

        self.lbl_pdf = tk.Label(self, text="Nenhum PDF selecionado", anchor="w", justify="left", wraplength=500)
        self.lbl_pdf.grid(row=0, column=1, sticky="w", **padding)

        btn_template = tk.Button(self, text="Selecionar Modelo Word", command=self.select_template)
        btn_template.grid(row=1, column=0, sticky="w", **padding)

        self.lbl_template = tk.Label(self, text="Nenhum modelo selecionado", anchor="w", wraplength=500)
        self.lbl_template.grid(row=1, column=1, sticky="w", **padding)

        tk.Label(self, text="Itens desejados (ex: 3,4,10)").grid(row=2, column=0, sticky="w", **padding)
        self.entry_itens = tk.Entry(self, width=50)
        self.entry_itens.grid(row=2, column=1, sticky="w", **padding)

        tk.Label(self, text="Nome da empresa").grid(row=3, column=0, sticky="w", **padding)
        self.entry_nome = tk.Entry(self, width=50)
        self.entry_nome.grid(row=3, column=1, sticky="w", **padding)

        tk.Label(self, text="CNPJ").grid(row=4, column=0, sticky="w", **padding)
        self.entry_cnpj = tk.Entry(self, width=50)
        self.entry_cnpj.grid(row=4, column=1, sticky="w", **padding)

        tk.Label(self, text="Nome do representante").grid(row=5, column=0, sticky="w", **padding)
        self.entry_representante = tk.Entry(self, width=50)
        self.entry_representante.grid(row=5, column=1, sticky="w", **padding)

        tk.Label(self, text="Nº do ofício").grid(row=6, column=0, sticky="w", **padding)
        self.entry_oficio = tk.Entry(self, width=50)
        self.entry_oficio.grid(row=6, column=1, sticky="w", **padding)

        tk.Label(self, text="E-mail").grid(row=7, column=0, sticky="w", **padding)
        self.entry_email = tk.Entry(self, width=50)
        self.entry_email.grid(row=7, column=1, sticky="w", **padding)

        btn_generate = tk.Button(self, text="Gerar Ofício", command=self.generate_office)
        btn_generate.grid(row=8, column=0, columnspan=2, pady=20)

    def select_pdfs(self) -> None:
        paths = filedialog.askopenfilenames(
            title="Selecione os PDFs da ATA",
            filetypes=(("Arquivos PDF", "*.pdf"), ("Todos os arquivos", "*.*")),
        )
        if paths:
            self.pdf_paths = list(paths)
            filenames = ", ".join(os.path.basename(path) for path in self.pdf_paths)
            self.lbl_pdf.config(text=filenames)
        else:
            self.pdf_paths = []
            self.lbl_pdf.config(text="Nenhum PDF selecionado")

    def select_template(self) -> None:
        path = filedialog.askopenfilename(
            title="Selecione o modelo Word",
            filetypes=(("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")),
        )
        if path:
            self.template_path = path
            self.lbl_template.config(text=os.path.basename(path))
        else:
            self.template_path = None
            self.lbl_template.config(text="Nenhum modelo selecionado")

    def generate_office(self) -> None:
        if not self.pdf_paths:
            messagebox.showwarning("Atenção", "Selecione ao menos um PDF da ATA.")
            return
        if not self.template_path or not os.path.exists(self.template_path):
            messagebox.showwarning("Atenção", "Selecione um modelo Word válido.")
            return

        selected_tokens = parse_selected_items(self.entry_itens.get())
        all_items: List[Dict[str, str]] = []
        try:
            for path in self.pdf_paths:
                items = extract_items_from_pdf(path)
                all_items.extend(items)
        except RuntimeError as exc:
            messagebox.showerror("Erro", str(exc))
            return

        filtered_items = filter_items(all_items, selected_tokens)
        if not filtered_items:
            messagebox.showwarning("Atenção", "Nenhum item correspondente foi encontrado.")
            return

        try:
            document = Document(self.template_path)
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível abrir o modelo Word: {exc}")
            return

        replacements = {
            "nome": self.entry_nome.get().strip(),
            "empresa": self.entry_nome.get().strip(),
            "cnpj": self.entry_cnpj.get().strip(),
            "representante": self.entry_representante.get().strip(),
            "oficio": self.entry_oficio.get().strip(),
            "ofício": self.entry_oficio.get().strip(),
            "email": self.entry_email.get().strip(),
            "e-mail": self.entry_email.get().strip(),
        }
        replacements = {k: v for k, v in replacements.items() if v}
        if replacements:
            replace_placeholders(document, replacements)

        try:
            fill_items_table(document, filtered_items)
        except RuntimeError as exc:
            messagebox.showerror("Erro", str(exc))
            return

        empresa = self.entry_nome.get().strip() or "sem_empresa"
        safe_empresa = re.sub(r"[^A-Za-z0-9_-]", "_", empresa)[:40] or "empresa"
        filename = f"Oficio_Preenchido_{datetime.now():%Y%m%d_%H%M%S}_{safe_empresa}.docx"
        try:
            document.save(filename)
        except Exception as exc:
            messagebox.showerror("Erro", f"Falha ao salvar o arquivo Word: {exc}")
            return

        messagebox.showinfo("Sucesso", f"Ofício gerado com sucesso!\nArquivo salvo como: {filename}")


def main() -> None:
    app = AtaApp()
    app.mainloop()


if __name__ == "__main__":
    main()
